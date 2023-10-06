import sys
from docx import Document

commit_sha = sys.argv[1]
pylint_score = sys.argv[2]

document = Document()
document.add_heading('Pylint Report', 0)

document.add_paragraph(f'Commit SHA: {commit_sha}')
document.add_paragraph(f'Pylint Score: {pylint_score}')

document.save(f'pylint_report_{commit_sha}.docx')
